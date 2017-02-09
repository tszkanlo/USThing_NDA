import docx
import sys
from datetime import datetime

document = docx.Document('usthing_nda.docx')

for paragraph in document.paragraphs:
    words = paragraph.text.split()
    for i in range(0, len(words)):
        if words[i] == '[DATE]':
            words[i] = datetime.now().strftime('%d, %B %Y')
            paragraph.text = ' '.join(words)
        if words[i] == '[RECEIVING_PARTY],':
            words[i] = sys.argv[1]
            paragraph.text = ' '.join(words)

for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                style = paragraph.style
                words = paragraph.text.split()
                for i in range(0, len(words)):
                    if words[i] == '[DATE]':
                        words[i] = datetime.now().strftime('%d, %B %Y')
                        cell.text = ' '.join(words)
                    if words[i] == '[RECEIVING_PARTY]':
                        words[i] = sys.argv[1]
                        cell.text = ' '.join(words)
                paragraph.style = style

document.save('usthing_nda_'+sys.argv[1]+'.docx')
